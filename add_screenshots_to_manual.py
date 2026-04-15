#!/usr/bin/env python3
"""
Script auxiliar para insertar screenshots en el manual DOCX
Permite agregar imágenes de forma automática a partir de una carpeta
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

def get_image_files(directory):
    """Obtiene todas las imágenes de un directorio"""
    images = []
    supported_formats = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
    
    if not Path(directory).exists():
        print(f"❌ Directorio no encontrado: {directory}")
        return images
    
    for file in sorted(Path(directory).glob('*')):
        if file.suffix.lower() in supported_formats:
            images.append(file)
    
    return images

def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_image_after_section(doc, section_title, image_path, caption=""):
    """Inserta una imagen después de un título específico"""
    found = False

    for para in doc.paragraphs:
        if section_title in para.text and para.style.name.startswith('Heading'):
            image_para = insert_paragraph_after(para)
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            try:
                run = image_para.add_run()
                run.add_picture(str(image_path), width=Inches(6.0))

                if caption:
                    caption_para = insert_paragraph_after(image_para, caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                print(f"✅ Imagen insertada en '{section_title}'")
                found = True
                break
            except Exception as e:
                print(f"❌ Error al insertar imagen: {e}")
                return False

    if not found:
        print(f"⚠️  No se encontró la sección: '{section_title}'")

    return found

def add_screenshots_to_manual(manual_path, screenshots_dir, output_path=None):
    """Agrega screenshots al manual desde un directorio"""
    
    if output_path is None:
        output_path = manual_path.replace('.docx', '_con_screenshots.docx')
    
    print(f"📖 Abriendo manual: {manual_path}")
    doc = Document(manual_path)
    
    print(f"📸 Buscando screenshots en: {screenshots_dir}")
    images = get_image_files(screenshots_dir)
    
    if not images:
        print("❌ No se encontraron imágenes en el directorio especificado")
        return False
    
    print(f"📊 Se encontraron {len(images)} imágenes\n")
    
    # Mapeo de secciones a imágenes (asume nombres de archivo específicos)
    section_mappings = {
        '2. Acceso al Sistema': ['01-login-studio.png', 'login.png'],
        '3. Panel Principal': ['02-dashboard-cursos.png', 'dashboard.png'],
        '5. Plantillas de Prueba': ['04-banco-preguntas-real.png', 'banco-preguntas-real.png', 'question-bank.png', 'bank-questions.png'],
        '6. Plantillas de Curso': ['05-library-assets.png', 'library-assets.png', 'course-templates.png'],
        '8. Gestión de Lecciones': ['03-editor-leccion.png', 'editor-leccion.png', 'lesson-editor.png'],
        '10. Ejercicios': ['06-rubricas.png', 'rubricas.png', 'exercises.png'],
        '11. Configuración': ['settings.png', 'configuracion.png', 'options.png', 'ajustes.png'],
        '12. Panel de Administrador': ['07-admin-usuarios.png', 'admin-usuarios.png', 'admin-panel.png']
    }
    
    inserted_count = 0
    
    for section, possible_names in section_mappings.items():
        for image in images:
            image_lower = image.name.lower()
            if any(name.lower() in image_lower for name in possible_names):
                caption = f"Figura: {section.split('. ')[1]}"
                if insert_image_after_section(doc, section.split('. ')[1], image, caption):
                    inserted_count += 1
                    images.remove(image)
                break
    
    print(f"\n✅ {inserted_count} imágenes insertadas")
    print(f"⚠️  {len(images)} imágenes no utilizadas")
    
    print(f"\n💾 Guardando manual actualizado en: {output_path}")
    doc.save(output_path)
    
    print("\n✅ Proceso completado exitosamente")
    return True

def main():
    """Función principal"""
    
    if len(sys.argv) < 2:
        print("🔧 Script para insertar screenshots en manual DOCX\n")
        print("Uso:")
        print("  python3 add_screenshots_to_manual.py <directorio_screenshots>")
        print("  python3 add_screenshots_to_manual.py <directorio_screenshots> <manual_entrada.docx> <salida.docx>")
        print("\nEjemplo:")
        print("  python3 add_screenshots_to_manual.py ./screenshots")
        print("  python3 add_screenshots_to_manual.py ./screenshots Manual_Usuario_OpenCCB_Studio.docx Manual_Completo.docx")
        sys.exit(1)
    
    screenshots_dir = sys.argv[1]
    manual_input = sys.argv[2] if len(sys.argv) > 2 else "Manual_Usuario_OpenCCB_Studio.docx"
    manual_output = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not Path(manual_input).exists():
        print(f"❌ Manual no encontrado: {manual_input}")
        sys.exit(1)
    
    success = add_screenshots_to_manual(manual_input, screenshots_dir, manual_output)
    sys.exit(0 if success else 1)

if __name__ == '__main__':
    main()