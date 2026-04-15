#!/usr/bin/env python3
"""
Generador del Manual de Usuario OpenCCB - Studio
Crea un documento DOCX profesional con guías de usuario
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Establece el color de fondo de una celda"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """Agrega un encabezado con estilo"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_styled_paragraph(doc, text, bold=False, italic=False, color=None):
    """Agrega un párrafo con estilo opcional"""
    p = doc.add_paragraph(text)
    if bold or italic or color:
        for run in p.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = color
    return p

def create_manual():
    """Crea el manual de usuario completo"""
    
    doc = Document()
    
    # ============================================================================
    # PORTADA
    # ============================================================================
    title = doc.add_heading('MANUAL DE USUARIO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('OpenCCB Studio', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_date = doc.add_paragraph(f'Versión 1.0 - {datetime.now().strftime("%d de %B de %Y")}')
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Descripción
    description = doc.add_paragraph(
        'Manual completo para el uso de OpenCCB Studio, '
        'plataforma de gestión de cursos en línea con inteligencia artificial integrada.'
    )
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ============================================================================
    # TABLA DE CONTENIDOS
    # ============================================================================
    add_heading(doc, 'Tabla de Contenidos', level=1)
    
    toc_items = [
        '1. Introducción',
        '2. Acceso al Sistema',
        '3. Panel Principal',
        '4. Gestión de Usuarios',
        '5. Plantillas de Prueba (Test Templates)',
        '6. Plantillas de Curso',
        '7. Creación de Cursos',
        '8. Gestión de Lecciones',
        '9. Uso de IA en Contenidos',
        '10. Ejercicios y Evaluaciones',
        '11. Configuración de Organización',
        '12. Panel de Administrador',
        '13. Análisis y Reportes',
        '14. Resolución de Problemas'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # ============================================================================
    # 1. INTRODUCCIÓN
    # ============================================================================
    add_heading(doc, '1. Introducción', level=1)
    
    doc.add_paragraph(
        'OpenCCB Studio es una plataforma moderna de gestión de cursos en línea diseñada '
        'para instructores y administradores educativos. Proporciona herramientas potentes '
        'para crear, gestionar y entregar contenido educativo de alta calidad.'
    )
    
    add_heading(doc, 'Características Principales', level=2)
    
    features = [
        'Creación flexible de cursos desde plantillas',
        'Generación de contenido asistida por IA',
        'Gestión de ejercicios y evaluaciones',
        'Análisis detallado del progreso estudiantil',
        'Integración con múltiples herramientas educativas',
        'Interfaz intuitiva y responsive',
        'Soporte para múltiples organizaciones',
        'Control de acceso basado en roles'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================================================
    # 2. ACCESO AL SISTEMA
    # ============================================================================
    add_heading(doc, '2. Acceso al Sistema', level=1)
    
    add_heading(doc, 'Login', level=2)
    doc.add_paragraph(
        '1. Abra el navegador e ingrese a studio.norteamericano.com'
    )
    doc.add_paragraph(
        '2. Ingrese su email registrado en el campo "Email"'
    )
    doc.add_paragraph(
        '3. Ingrese su contraseña en el campo "Contraseña"'
    )
    doc.add_paragraph(
        '4. Haga clic en el botón "Iniciar Sesión"'
    )
    
    doc.add_paragraph(
        '\n✅ Consejo: Si olvida su contraseña, use la opción "¿Olvidó su contraseña?" '
        'en la página de login.'
    )
    
    add_heading(doc, 'Tipos de Usuarios', level=2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Tipo de Usuario'
    header_cells[1].text = 'Permisos'
    
    data = [
        ['Estudiante', 'Ver cursos, completar lecciones, realizar ejercicios, ver calificaciones'],
        ['Instructor', 'Crear cursos, gestionar contenido, calificar, ver análisis'],
        ['Administrador', 'Gestión completa del sistema, usuarios, organizaciones, configuración']
    ]
    
    for i, (role, permissions) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = role
        row_cells[1].text = permissions
    
    doc.add_page_break()
    
    # ============================================================================
    # 3. PANEL PRINCIPAL
    # ============================================================================
    add_heading(doc, '3. Panel Principal', level=1)
    
    doc.add_paragraph(
        'Después de iniciar sesión, accederá al panel principal que muestra '
        'un resumen de sus cursos y actividades recientes.'
    )
    
    add_heading(doc, 'Componentes del Panel', level=2)
    
    components = {
        'Barra de Navegación Superior': 
            'Acceso rápido a Cursos, Administración, Perfil y Configuración',
        'Mis Cursos': 
            'Lista de cursos en los que participa (como instructor o estudiante)',
        'Actividades Recientes': 
            'Muestra anuncios, tareas y eventos próximos',
        'Búsqueda y Filtros': 
            'Herramientas para encontrar cursos específicos',
        'Estadísticas Rápidas':
            'Resumen de estudiantes, tareas pendientes y calificaciones'
    }
    
    for component, description in components.items():
        p = doc.add_paragraph(f'{component}: {description}')
        p.style = 'List Bullet'
    
    doc.add_page_break()
    
    # ============================================================================
    # 4. GESTIÓN DE USUARIOS
    # ============================================================================
    add_heading(doc, '4. Gestión de Usuarios', level=1)
    
    add_heading(doc, 'Panel de Administración de Usuarios', level=2)
    
    doc.add_paragraph(
        '📍 Ubicación: Administración → Usuarios'
    )
    
    doc.add_paragraph(
        'En este panel puede gestionar todos los usuarios de su organización.'
    )
    
    add_heading(doc, 'Crear un Nuevo Usuario', level=3)
    
    steps = [
        'Haga clic en el botón "Nuevo Usuario"',
        'Complete los siguientes campos:',
        '  • Email: Correo único del usuario',
        '  • Nombre Completo: Nombre del usuario',
        '  • Contraseña: Al menos 8 caracteres',
        '  • Rol: Seleccione Estudiante, Instructor o Administrador',
        'Haga clic en "Crear Usuario"',
        '✅ El usuario recibirá un email de confirmación'
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number' if not step.startswith('  ') else 'List Bullet')
    
    add_heading(doc, 'Editar Usuario Existente', level=3)
    
    doc.add_paragraph(
        '1. Busque el usuario en la lista o use el buscador'
    )
    doc.add_paragraph(
        '2. Haga clic en el icono "Editar" (lápiz)'
    )
    doc.add_paragraph(
        '3. Modifique los campos necesarios'
    )
    doc.add_paragraph(
        '4. Haga clic en "Guardar Cambios"'
    )
    
    doc.add_page_break()
    
    # ============================================================================
    # 5. PLANTILLAS DE PRUEBA
    # ============================================================================
    add_heading(doc, '5. Plantillas de Prueba (Test Templates)', level=1)
    
    doc.add_paragraph(
        'Las Plantillas de Prueba son colecciones predefinidas de preguntas que puede '
        'usar para crear evaluaciones rápidamente. OpenCCB proporciona plantillas listas '
        'para usar o puede crear personalizadas.'
    )
    
    add_heading(doc, 'Acceder a Plantillas de Prueba', level=2)
    
    doc.add_paragraph(
        '📍 Ubicación: Menú Principal → Plantillas de Prueba'
    )
    
    add_heading(doc, 'Crear una Nueva Plantilla de Prueba', level=3)
    
    doc.add_paragraph('1. Haga clic en "Crear Nueva Plantilla"')
    doc.add_paragraph('2. Complete la información básica:')
    doc.add_paragraph('   • Nombre: "Prueba Inicial de Matemáticas"', style='List Bullet')
    doc.add_paragraph('   • Descripción: Breve resumen del contenido', style='List Bullet')
    doc.add_paragraph('   • Categoría: Seleccione una categoría existente', style='List Bullet')
    doc.add_paragraph('3. Haga clic en "Agregar Preguntas"')
    
    add_heading(doc, 'Agregar Preguntas a la Plantilla', level=3)
    
    question_types = [
        'Opción Múltiple: El estudiante selecciona una respuesta correcta',
        'Verdadero/Falso: Afirmación para validar',
        'Respuesta Corta: Campo de texto para respuesta breve',
        'Ensayo: Espacio libre para respuestas largas',
        'Emparejamiento: Vincular elementos de dos listas'
    ]
    
    doc.add_paragraph('Tipos de preguntas disponibles:', style='List Bullet')
    for qtype in question_types:
        doc.add_paragraph(qtype, style='List Bullet')
    
    doc.add_paragraph(
        '💡 Sugerencia: Use la IA para generar preguntas automáticamente '
        '(ver sección de IA).'
    )
    
    doc.add_page_break()
    
    # ============================================================================
    # 6. PLANTILLAS DE CURSO
    # ============================================================================
    add_heading(doc, '6. Plantillas de Curso', level=1)
    
    doc.add_paragraph(
        'Las Plantillas de Curso son estructuras predefinidas para cursos con '
        'módulos, lecciones y evaluaciones ya configurados. Permiten crear cursos '
        'consistentes rápidamente.'
    )
    
    add_heading(doc, 'Ver Plantillas de Curso', level=2)
    
    doc.add_paragraph(
        '📍 Ubicación: Menú Principal → Plantillas de Curso'
    )
    
    add_heading(doc, 'Crear una Nueva Plantilla de Curso', level=3)
    
    doc.add_paragraph('1. Haga clic en "Crear Nueva Plantilla"')
    doc.add_paragraph('2. Complete los detalles:')
    doc.add_paragraph('   • Nombre: "Desarrollo Web Avanzado"', style='List Bullet')
    doc.add_paragraph('   • Descripción detallada', style='List Bullet')
    doc.add_paragraph('   • Duración esperada: número de semanas', style='List Bullet')
    doc.add_paragraph('   • Nivel de dificultad: Básico/Intermedio/Avanzado', style='List Bullet')
    
    doc.add_paragraph('3. Configure la estructura del curso:')
    doc.add_paragraph('   • Módulo 1: Fundamentos', style='List Bullet')
    doc.add_paragraph('   • Módulo 2: Proyectos Prácticos', style='List Bullet')
    doc.add_paragraph('   • Módulo 3: Evaluación Final', style='List Bullet')
    
    doc.add_paragraph('4. Para cada módulo, agregue lecciones:')
    doc.add_paragraph('   • Título', style='List Bullet')
    doc.add_paragraph('   • Descripción', style='List Bullet')
    doc.add_paragraph('   • Contenido (texto, videos, recursos)', style='List Bullet')
    doc.add_paragraph('   • Evaluación asociada', style='List Bullet')
    
    doc.add_paragraph('5. Configure evaluaciones:')
    doc.add_paragraph('   • Tipo: Prueba, Tarea, Proyecto', style='List Bullet')
    doc.add_paragraph('   • Rubrica de calificación', style='List Bullet')
    doc.add_paragraph('   • Fecha de vencimiento', style='List Bullet')
    
    doc.add_paragraph('6. Haga clic en "Guardar Plantilla"')
    
    doc.add_page_break()
    
    # ============================================================================
    # 7. CREACIÓN DE CURSOS
    # ============================================================================
    add_heading(doc, '7. Creación de Cursos - Flujo Completo', level=1)
    
    doc.add_paragraph(
        'Este es el proceso completo para crear un curso desde cero usando plantillas. '
        'El flujo recomendado es: Plantilla de Prueba → Plantilla de Curso → Curso Final'
    )
    
    add_heading(doc, 'Paso 1: Crear Plantilla de Prueba', level=2)
    
    doc.add_paragraph(
        'Primero, prepare las evaluaciones que usará en su curso:'
    )
    doc.add_paragraph(
        '1. Vaya a "Plantillas de Prueba" y cree 2-3 plantillas para su contenido'
    )
    doc.add_paragraph(
        '2. Agregue preguntas representativas del contenido del curso'
    )
    doc.add_paragraph(
        '3. Pruebe las plantillas para asegurar claridad (use IA para mejorar preguntas)'
    )
    
    add_heading(doc, 'Paso 2: Crear Plantilla de Curso', level=2)
    
    doc.add_paragraph(
        'Luego, cree la estructura base del curso:'
    )
    doc.add_paragraph(
        '1. Vaya a "Plantillas de Curso" y haga clic en "Crear Nueva Plantilla"'
    )
    doc.add_paragraph(
        '2. Defina la estructura con módulos y lecciones'
    )
    doc.add_paragraph(
        '3. Asocie las plantillas de prueba con las lecciones correspondientes'
    )
    doc.add_paragraph(
        '4. Agregue recursos, materiales complementarios'
    )
    doc.add_paragraph(
        '5. Guarde la plantilla para reutilizarla en futuro'
    )
    
    add_heading(doc, 'Paso 3: Crear el Curso Final', level=2)
    
    doc.add_paragraph(
        'Finalmente, cree el curso basado en la plantilla:'
    )
    doc.add_paragraph(
        '📍 Ubicación: Menú Principal → Mis Cursos → Crear Curso'
    )
    
    doc.add_paragraph('\n1. Seleccione una Plantilla de Curso')
    doc.add_paragraph('2. Complete la información del curso:')
    doc.add_paragraph('   • Nombre del Curso', style='List Bullet')
    doc.add_paragraph('   • Descripción (pública para estudiantes)', style='List Bullet')
    doc.add_paragraph('   • Código del Curso (ej: MAT-101)', style='List Bullet')
    doc.add_paragraph('   • Instructor Principal', style='List Bullet')
    doc.add_paragraph('   • Horario y Ubicación (si aplica)', style='List Bullet')
    doc.add_paragraph('   • Capacidad Máxima de Estudiantes', style='List Bullet')
    
    doc.add_paragraph('3. Configure acceso y visibilidad:')
    doc.add_paragraph('   • Privado / Público', style='List Bullet')
    doc.add_paragraph('   • Requiere Código de Inscripción', style='List Bullet')
    doc.add_paragraph('   • Aprobación Manual de Estudiantes', style='List Bullet')
    
    doc.add_paragraph('4. Establezca fechas importantes:')
    doc.add_paragraph('   • Fecha de Inicio', style='List Bullet')
    doc.add_paragraph('   • Fecha de Fin', style='List Bullet')
    doc.add_paragraph('   • Fechas de Evaluaciones', style='List Bullet')
    
    doc.add_paragraph('5. Agregue co-instructores (opcional)')
    doc.add_paragraph('6. Haga clic en "Crear Curso"')
    
    doc.add_paragraph('\n✅ Su curso está listo. Ahora puede published o invitar estudiantes.')
    
    doc.add_page_break()
    
    # ============================================================================
    # 8. GESTIÓN DE LECCIONES
    # ============================================================================
    add_heading(doc, '8. Gestión de Lecciones', level=1)
    
    add_heading(doc, 'Editor de Lecciones', level=2)
    
    doc.add_paragraph(
        'Dentro de un curso, puede editar cada lección de forma detallada.'
    )
    
    doc.add_paragraph(
        '📍 Ubicación: Curso → Lecciones → Seleccionar Lección'
    )
    
    add_heading(doc, 'Componentes de una Lección', level=3)
    
    components = {
        'Título y Descripción': 'Información básica de la lección',
        'Contenido de Texto': 'Texto formateado, párrafos, listas',
        'Videos': 'Embeber videos de YouTube u otros proveedores',
        'Archivos': 'Documentos PDF, presentaciones, recursos',
        'Actividades Interactivas': 'Ejercicios, tareas, discusiones',
        'Quiz': 'Evaluaciones integradas usando plantillas de prueba',
        'Foro de Discusión': 'Espacio para preguntas y reflexión',
        'Recursos Externa': 'Enlaces a sitios web, herramientas'
    }
    
    for component, desc in components.items():
        p = doc.add_paragraph(f'{component}: {desc}')
        p.style = 'List Bullet'
    
    add_heading(doc, 'Agregar Contenido a una Lección', level=3)
    
    doc.add_paragraph('1. Abra la lección para editar')
    doc.add_paragraph('2. Haga clic en "Agregar Contenido"')
    doc.add_paragraph('3. Seleccione el tipo de contenido (Texto, Video, Archivo, etc.)')
    doc.add_paragraph('4. Complete los detalles')
    doc.add_paragraph('5. Haga clic en "Guardar"')
    
    doc.add_paragraph(
        '💡 Sugerencia: Use el editor visual con arrastrar-soltar para organizar '
        'bloques de contenido fácilmente.'
    )
    
    doc.add_page_break()
    
    # ============================================================================
    # 9. USO DE IA EN CONTENIDOS
    # ============================================================================
    add_heading(doc, '9. Uso de IA en Contenidos', level=1)
    
    doc.add_paragraph(
        'OpenCCB integra inteligencia artificial para ayudar en la creación de contenido, '
        'generación de preguntas, mejora de textos y más.'
    )
    
    add_heading(doc, 'Características de IA Disponibles', level=2)
    
    ia_features = {
        'Generación de Preguntas': 
            'Cree automáticamente preguntas basadas en un tema o contenido',
        'Mejora de Textos': 
            'Mejore redacción, gramática y claridad de contenido',
        'Resúmenes Automáticos': 
            'Cree resúmenes de textos extensos',
        'Sugerencias de Actividades': 
            'Propuestas de ejercicios para reforzar conceptos',
        'Análisis de Respuestas': 
            'Evaluación automática de respuestas de ensayo',
        'Generación de Rubricas': 
            'Cree rúbricas de evaluación automáticamente',
        'Explicaciones Personalizadas':
            'Genere explicaciones adaptadas al nivel del estudiante'
    }
    
    for feature, description in ia_features.items():
        p = doc.add_paragraph(f'{feature}: {description}')
        p.style = 'List Bullet'
    
    add_heading(doc, 'Cómo Usar IA para Generar Preguntas', level=3)
    
    doc.add_paragraph('1. En una Plantilla de Prueba, haga clic en "Generar con IA"')
    doc.add_paragraph('2. Ingrese el tema: "Teoría de la Relatividad"')
    doc.add_paragraph('3. Seleccione parámetros:')
    doc.add_paragraph('   • Cantidad de preguntas: 5-10', style='List Bullet')
    doc.add_paragraph('   • Tipo: Opción Múltiple, Verdadero/Falso, etc.', style='List Bullet')
    doc.add_paragraph('   • Nivel de dificultad', style='List Bullet')
    doc.add_paragraph('4. Haga clic en "Generar"')
    doc.add_paragraph('5. Revise y edite las preguntas generadas')
    doc.add_paragraph('6. Haga clic en "Aceptar" para agregar al banco de preguntas')
    
    add_heading(doc, 'Cómo Usar IA para Mejorar Contenido', level=3)
    
    doc.add_paragraph('1. En el editor de lecciones, seleccione un bloque de texto')
    doc.add_paragraph('2. Haga clic en el botón "✨ Mejorar con IA"')
    doc.add_paragraph('3. Seleccione el tipo de mejora:')
    doc.add_paragraph('   • Mejorar redacción y claridad', style='List Bullet')
    doc.add_paragraph('   • Hacer más conciso', style='List Bullet')
    doc.add_paragraph('   • Expandir con más detalles', style='List Bullet')
    doc.add_paragraph('   • Traducir a otro idioma', style='List Bullet')
    doc.add_paragraph('4. Revise la propuesta y acepte o rechace')
    
    add_heading(doc, 'Análisis de Respuestas con IA', level=3)
    
    doc.add_paragraph(
        'Al habilitar calificación automática en una evaluación:'
    )
    doc.add_paragraph('1. La IA analiza respuestas de ensayo automáticamente')
    doc.add_paragraph('2. Proporciona una calificación inicial basada en la rúbrica')
    doc.add_paragraph('3. Usted puede revisar y ajustar calificaciones')
    doc.add_paragraph('4. Se proporciona feedback personalizado al estudiante')
    
    doc.add_page_break()
    
    # ============================================================================
    # 10. EJERCICIOS Y EVALUACIONES
    # ============================================================================
    add_heading(doc, '10. Ejercicios y Evaluaciones', level=1)
    
    doc.add_paragraph(
        'Los ejercicios y evaluaciones son componentes críticos para medir '
        'el aprendizaje de los estudiantes.'
    )
    
    add_heading(doc, 'Tipos de Evaluación', level=2)
    
    types = {
        'Quiz/Prueba': 
            'Preguntas automáticamente calificadas (opción múltiple, verdadero/falso)',
        'Tarea': 
            'Trabajo que requiere revisión manual del instructor',
        'Proyecto': 
            'Trabajo colaborativo o individual de duración más larga',
        'Discusión': 
            'Participación en foros sobre temas específicos',
        'Participación':
            'Asistencia y participación en clase (si hay sesiones sincrónicas)'
    }
    
    for eval_type, desc in types.items():
        p = doc.add_paragraph(f'{eval_type}: {desc}')
        p.style = 'List Bullet'
    
    add_heading(doc, 'Crear una Evaluación', level=3)
    
    doc.add_paragraph('1. En el curso, vaya a "Evaluaciones"')
    doc.add_paragraph('2. Haga clic en "Crear Nueva Evaluación"')
    doc.add_paragraph('3. Complete información básica:')
    doc.add_paragraph('   • Nombre: "Quiz Unidad 1"', style='List Bullet')
    doc.add_paragraph('   • Descripción', style='List Bullet')
    doc.add_paragraph('   • Tipo: Prueba, Tarea, Proyecto', style='List Bullet')
    doc.add_paragraph('   • Peso en calificación: %', style='List Bullet')
    
    doc.add_paragraph('4. Para Pruebas:')
    doc.add_paragraph('   • Seleccione plantilla de prueba o cree nueva', style='List Bullet')
    doc.add_paragraph('   • Configure tiempo límite', style='List Bullet')
    doc.add_paragraph('   • Número de intentos permitidos', style='List Bullet')
    doc.add_paragraph('   • Mostrar respuestas después de enviar', style='List Bullet')
    
    doc.add_paragraph('5. Para Tareas/Proyectos:')
    doc.add_paragraph('   • Fecha de vencimiento', style='List Bullet')
    doc.add_paragraph('   • Rúbrica de evaluación', style='List Bullet')
    doc.add_paragraph('   • Archivos permitidos', style='List Bullet')
    doc.add_paragraph('   • Límite de tamaño', style='List Bullet')
    
    doc.add_paragraph('6. Haga clic en "Crear"')
    
    add_heading(doc, 'Calificar Evaluaciones', level=3)
    
    doc.add_paragraph('📍 Ubicación: Curso → Calificaciones')
    
    doc.add_paragraph('1. Seleccione la evaluación a calificar')
    doc.add_paragraph('2. Vea lista de envíos de estudiantes')
    doc.add_paragraph('3. Haga clic en un envío para revisar')
    doc.add_paragraph('4. Agregue comentarios y calificación')
    doc.add_paragraph('5. (Opcional) Adjunte una rúbrica detallada')
    doc.add_paragraph('6. Haga clic en "Guardar Calificación" para notificar al estudiante')
    
    doc.add_page_break()
    
    # ============================================================================
    # 11. CONFIGURACIÓN DE ORGANIZACIÓN
    # ============================================================================
    add_heading(doc, '11. Configuración de Organización', level=1)
    
    doc.add_paragraph(
        '📍 Ubicación: Configuración → Organización'
    )
    
    add_heading(doc, 'Información Básica', level=2)
    
    doc.add_paragraph('Configure detalles de su organización:')
    doc.add_paragraph('• Nombre oficial', style='List Bullet')
    doc.add_paragraph('• Logo (aparece en todas las páginas)', style='List Bullet')
    doc.add_paragraph('• Favicon (ícono en pestaña)', style='List Bullet')
    doc.add_paragraph('• Descripción', style='List Bullet')
    doc.add_paragraph('• Sitio web', style='List Bullet')
    
    add_heading(doc, 'Configuración de Branding', level=2)
    
    doc.add_paragraph('Personalice la apariencia visual:')
    doc.add_paragraph('• Color principal', style='List Bullet')
    doc.add_paragraph('• Fuente tipográfica', style='List Bullet')
    doc.add_paragraph('• Tema claro/oscuro', style='List Bullet')
    
    add_heading(doc, 'Servicios de Email', level=2)
    
    doc.add_paragraph('Configure cómo la plataforma envía correos:')
    doc.add_paragraph('• Servidor SMTP', style='List Bullet')
    doc.add_paragraph('• Email de remitente', style='List Bullet')
    doc.add_paragraph('• Credenciales de autenticación', style='List Bullet')
    doc.add_paragraph('• Templates de email personalizados', style='List Bullet')
    
    add_heading(doc, 'Plantillas de Email Personalizadas', level=3)
    
    doc.add_paragraph(
        'Personalice los emails enviados por el sistema:'
    )
    doc.add_paragraph('1. Vaya a "Configuración → Plantillas de Email"')
    doc.add_paragraph('2. Seleccione el tipo de email a personalizar:')
    doc.add_paragraph('   • Bienvenida de estudiante', style='List Bullet')
    doc.add_paragraph('   • Respuesta en foro', style='List Bullet')
    doc.add_paragraph('   • Notificación de calificación', style='List Bullet')
    doc.add_paragraph('   • Recordatorio de fecha límite', style='List Bullet')
    doc.add_paragraph('3. Edite el asunto y contenido')
    doc.add_paragraph('4. Use variables como {{estudante_nombre}}, {{curso_nombre}}, etc.')
    doc.add_paragraph('5. Vista previa y guarde')
    
    doc.add_page_break()
    
    # ============================================================================
    # 12. PANEL DE ADMINISTRADOR
    # ============================================================================
    add_heading(doc, '12. Panel de Administrador', level=1)
    
    doc.add_paragraph(
        '📍 Ubicación: Menú Principal → Administración'
    )
    
    doc.add_paragraph(
        'El panel de administrador proporciona herramientas para gestionar '
        'la plataforma a nivel de organización.'
    )
    
    add_heading(doc, 'Secciones Principales', level=2)
    
    sections = {
        'Usuarios': 
            'Crear, editar, eliminar usuarios y gestionar roles',
        'Cursos': 
            'Ver todos los cursos, gestionar inscripciones',
        'Reportes': 
            'Análisis de uso, progreso estudiantil, métricas',
        'Auditoría': 
            'Registro de cambios y actividades de usuarios',
        'Configuración':
            'Configuración general del sistema',
        'Uso de IA':
            'Monitorear uso de servicios de IA y límites'
    }
    
    for section, desc in sections.items():
        p = doc.add_paragraph(f'{section}: {desc}')
        p.style = 'List Bullet'
    
    add_heading(doc, 'Gestión de Usuarios Admin', level=3)
    
    doc.add_paragraph('1. Vaya a "Administración → Usuarios"')
    doc.add_paragraph('2. Puede:')
    doc.add_paragraph('   • Ver lista de todos los usuarios', style='List Bullet')
    doc.add_paragraph('   • Buscar por email o nombre', style='List Bullet')
    doc.add_paragraph('   • Crear nuevos usuarios', style='List Bullet')
    doc.add_paragraph('   • Editar información y roles', style='List Bullet')
    doc.add_paragraph('   • Desactivar/reactivar cuentas', style='List Bullet')
    doc.add_paragraph('   • Ver historial de actividad', style='List Bullet')
    
    add_heading(doc, 'Reportes y Análisis', level=3)
    
    doc.add_paragraph('Acceda a reportes detallados:')
    doc.add_paragraph(
        '📍 Ubicación: Administración → Reportes'
    )
    
    doc.add_paragraph('\nTipos de reportes disponibles:')
    doc.add_paragraph('• Inscripciones por curso', style='List Bullet')
    doc.add_paragraph('• Progreso estudiantil', style='List Bullet')
    doc.add_paragraph('• Uso de plataforma por usuario', style='List Bullet')
    doc.add_paragraph('• Asignaciones de calificaciones', style='List Bullet')
    doc.add_paragraph('• Actividad de foros', style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================================================
    # 13. ANÁLISIS Y REPORTES
    # ============================================================================
    add_heading(doc, '13. Análisis y Reportes', level=1)
    
    add_heading(doc, 'Dashboard del Instructor', level=2)
    
    doc.add_paragraph(
        'Cada instructor puede ver análisis detallado de sus cursos:'
    )
    
    doc.add_paragraph(
        '📍 Ubicación: Curso → Análisis'
    )
    
    doc.add_paragraph('\nMétricas disponibles:')
    doc.add_paragraph('• Promedio de calificaciones', style='List Bullet')
    doc.add_paragraph('• Tasa de finalización de lecciones', style='List Bullet')
    doc.add_paragraph('• Participación en foros', style='List Bullet')
    doc.add_paragraph('• Distribución de calificaciones', style='List Bullet')
    doc.add_paragraph('• Estudiantes en riesgo (bajo desempeño)', style='List Bullet')
    doc.add_paragraph('• Recursos más utilizados', style='List Bullet')
    
    add_heading(doc, 'Análisis Estudiantil Individual', level=3)
    
    doc.add_paragraph('1. En el curso, vaya a "Estudiantes"')
    doc.add_paragraph('2. Haga clic en un estudiante')
    doc.add_paragraph('3. Verá:')
    doc.add_paragraph('   • Calificaciones en todas las evaluaciones', style='List Bullet')
    doc.add_paragraph('   • Porcentaje de finalización de lecciones', style='List Bullet')
    doc.add_paragraph('   • Última actividad', style='List Bullet')
    doc.add_paragraph('   • Tiempo dedicado en la plataforma', style='List Bullet')
    doc.add_paragraph('   • Participación en foros', style='List Bullet')
    
    add_heading(doc, 'Exportar Reportes', level=3)
    
    doc.add_paragraph('Los reportes pueden exportarse en:')
    doc.add_paragraph('• PDF', style='List Bullet')
    doc.add_paragraph('• CSV (para Excel)', style='List Bullet')
    doc.add_paragraph('• JSON (para integración)', style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================================================
    # 14. RESOLUCIÓN DE PROBLEMAS
    # ============================================================================
    add_heading(doc, '14. Resolución de Problemas', level=1)
    
    add_heading(doc, 'Problemas Comunes', level=2)
    
    problems = [
        {
            'issue': '❌ No puedo iniciar sesión',
            'solutions': [
                'Verifique que escribió correctamente su email',
                'Haga clic en "¿Olvidó su contraseña?" para resetearla',
                'Limpie cookies y cache del navegador',
                'Intente en otro navegador (Chrome, Firefox, Safari)',
                'Asegúrese de que su cuenta fue activada'
            ]
        },
        {
            'issue': '❌ No puedo subir archivos',
            'solutions': [
                'Verifique el tamaño del archivo (máximo 100 MB)',
                'Confirme que el formato es permitido (PDF, DOCX, PPT, etc.)',
                'Intente subir archivos más pequeños primero',
                'Asegúrese de tener permisos suficientes',
                'Contacte a administrador si persiste el problema'
            ]
        },
        {
            'issue': '❌ La IA no genera preguntas',
            'solutions': [
                'Verifique que tenga saldo de uso de IA disponible',
                'Intente con un tema más específico',
                'Revise que el contenido no sea demasiado vago',
                'Espere un momento y reintente',
                'Contacte a soporte técnico'
            ]
        },
        {
            'issue': '❌ Los estudiantes no ven el curso',
            'solutions': [
                'Confirme que el curso está publicado (no en borrador)',
                'Verifique que el curso no está con acceso restringido',
                'Asegúrese de que los estudiantes están inscritos',
                'Revise el código de inscripción si es requerido',
                'Intente acceder como estudiante para probar'
            ]
        }
    ]
    
    for prob in problems:
        add_heading(doc, prob['issue'], level=3)
        for solution in prob['solutions']:
            doc.add_paragraph(solution, style='List Bullet')
        doc.add_paragraph()
    
    add_heading(doc, 'Contactar Soporte', level=2)
    
    doc.add_paragraph(
        'Si el problema persiste, contacte al equipo de soporte:'
    )
    doc.add_paragraph('Email: soporte@openccb.com', style='List Bullet')
    doc.add_paragraph('Teléfono: +1-XXX-XXX-XXXX', style='List Bullet')
    doc.add_paragraph('Chat en vivo: Disponible en horario comercial', style='List Bullet')
    
    doc.add_paragraph(
        '\n📋 Proporcione detalles precisos del problema:'
    )
    doc.add_paragraph('• Descripción paso a paso', style='List Bullet')
    doc.add_paragraph('• Captura de pantalla del error', style='List Bullet')
    doc.add_paragraph('• Navegador utilizado', style='List Bullet')
    doc.add_paragraph('• Hora del problema', style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================================================
    # APÉNDICE
    # ============================================================================
    add_heading(doc, 'Apéndice: Atajos de Teclado', level=1)
    
    keyboard_table = doc.add_table(rows=8, cols=2)
    keyboard_table.style = 'Light Grid Accent 1'
    
    header = keyboard_table.rows[0].cells
    header[0].text = 'Atajo'
    header[1].text = 'Función'
    
    shortcuts = [
        ['Ctrl + S / Cmd + S', 'Guardar página o formulario'],
        ['Ctrl + F / Cmd + F', 'Buscar en la página'],
        ['Tab', 'Navegar entre elementos'],
        ['Enter', 'Enviar formulario'],
        ['Esc', 'Cerrar diálogos'],
        ['Ctrl + Z / Cmd + Z', 'Deshacer últimas acciones'],
        ['Alt + ← / Cmd + ←', 'Atrás en navegador']
    ]
    
    for i, (shortcut, function) in enumerate(shortcuts, 1):
        row = keyboard_table.rows[i].cells
        row[0].text = shortcut
        row[1].text = function
    
    doc.add_paragraph()
    add_heading(doc, 'Glosario de Términos', level=1)
    
    terms = {
        'LMS': 'Learning Management System (Sistema de Gestión de Aprendizaje)',
        'SCORM': 'Estándar para contenido electrónico reutilizable',
        'Rúbrica': 'Criterios de evaluación con escalas de puntuación',
        'Widget': 'Componente pequeño embebible en lecciones',
        'API': 'Interface de Programación de Aplicaciones',
        'SSL': 'Protocolo de seguridad para conexiones web'
    }
    
    glossary_list = doc.add_paragraph()
    for term, definition in terms.items():
        p = doc.add_paragraph(f'{term}: {definition}')
        p.style = 'List Bullet'
    
    # ============================================================================
    # ÚLTIMA PÁGINA
    # ============================================================================
    doc.add_page_break()
    
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_text = final.add_run(
        '📚 FIN DEL MANUAL\n\n'
        'Para más información y actualizaciones, visite:\n'
        'www.openccb.com/docs\n\n'
        '© 2024 OpenCCB. Todos los derechos reservados.'
    )
    final_text.font.size = Pt(11)
    
    return doc

def main():
    """Función principal para generar el documento"""
    print("🚀 Generando Manual de Usuario OpenCCB Studio...")
    
    doc = create_manual()
    
    output_path = '/home/juan/dev/openccb/Manual_Usuario_OpenCCB_Studio.docx'
    doc.save(output_path)
    
    print(f"✅ Manual generado exitosamente en: {output_path}")
    print(f"📄 Archivo: Manual_Usuario_OpenCCB_Studio.docx")
    print(f"📊 Páginas: ~40-50 (sin screenshots)")
    print("\n💡 Próximos pasos:")
    print("1. Abra el documento en Microsoft Word o LibreOffice")
    print("2. Agregue screenshots del sistema en las secciones indicadas")
    print("3. Personalice los emails y datos de contacto")
    print("4. Exporte como PDF para distribución")

if __name__ == '__main__':
    main()